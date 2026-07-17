"""Local documents connector — indexes .md / .pdf / .docx (and plain text).

Useful for notes, papers, assignments, and any document folder you want to
search and chat over. Extracted text is chunked, embedded, and stored exactly
like the other connectors; the connector simply owns its own repo identity so
document content stays separable from code in the index.
"""
from __future__ import annotations

import os
from datetime import datetime
from pathlib import Path
from typing import Iterator

from connectors.base import BaseConnector, Document

DOC_EXTENSIONS = {".md", ".markdown", ".txt", ".rst", ".pdf", ".docx"}
DEFAULT_SKIP_DIRS = {".git", "node_modules", "__pycache__", ".venv", "venv", "dist", "build"}


class DocumentsConnector(BaseConnector):
    name = "documents"

    def __init__(self, config: dict):
        super().__init__(config)
        self.paths = self._as_list(config.get("paths"))
        self.recursive = bool(config.get("recursive", True))
        self.label = (config.get("label") or "documents").strip() or "documents"
        self.skip_dirs = set(config.get("skip_dirs", list(DEFAULT_SKIP_DIRS)))
        self.max_size_bytes = int(config.get("max_file_size_kb", 10_000)) * 1024
        self.min_text_chars = int(config.get("min_text_chars", 100))
        # Stored source with its own repo identity in the index.
        self.repo_name = self.label

    def get_repo_name(self) -> str:
        return self.repo_name

    def _as_list(self, value) -> list[str]:
        if not value:
            return []
        if isinstance(value, str):
            return [v.strip() for v in value.splitlines() if v.strip()]
        return [str(v).strip() for v in value if str(v).strip()]

    def _read(self, path: Path) -> str | None:
        ext = path.suffix.lower()
        try:
            if ext == ".pdf":
                return self._read_pdf(path)
            if ext == ".docx":
                return self._read_docx(path)
            try:
                return path.read_text(encoding="utf-8")
            except UnicodeDecodeError:
                return path.read_text(encoding="latin-1")
        except Exception:
            return None

    def _read_pdf(self, path: Path) -> str | None:
        try:
            from pypdf import PdfReader
        except ImportError:
            return None
        try:
            reader = PdfReader(str(path))
            return "\n\n".join(p.extract_text() or "" for p in reader.pages).strip()
        except Exception:
            return None

    def _read_docx(self, path: Path) -> str | None:
        try:
            from docx import Document as DocxDocument
        except ImportError:
            return None
        try:
            doc = DocxDocument(str(path))
            return "\n".join(p.text for p in doc.paragraphs if p.text).strip()
        except Exception:
            return None

    def load_documents(self) -> Iterator[Document]:
        if not self.paths:
            raise ValueError("Documents connector needs at least one path")

        roots: list[Path] = []
        for raw in self.paths:
            p = Path(raw)
            if p.is_file() or p.is_dir():
                roots.append(p)
        if not roots:
            raise ValueError("No valid documents path found")

        seen: set[str] = set()
        for root in roots:
            base = root if root.is_dir() else root.parent
            if root.is_file():
                candidates = [root]
            else:
                candidates = root.rglob("*") if self.recursive else root.glob("*")
            for fpath in candidates:
                if not fpath.is_file():
                    continue
                ext = fpath.suffix.lower()
                if ext not in DOC_EXTENSIONS:
                    continue
                if any(part in self.skip_dirs for part in fpath.parts):
                    continue
                try:
                    if fpath.stat().st_size > self.max_size_bytes:
                        continue
                except OSError:
                    continue
                try:
                    rel = str(fpath.relative_to(base)).replace("\\", "/")
                except ValueError:
                    rel = fpath.name
                if rel in seen:
                    continue
                seen.add(rel)

                content = self._read(fpath)
                if not content or not content.strip():
                    continue
                if len(content) < self.min_text_chars:
                    continue
                try:
                    mtime = datetime.fromtimestamp(fpath.stat().st_mtime)
                except OSError:
                    mtime = None

                yield Document(
                    id=f"doc:{self.label}:{rel}",
                    content=content,
                    title=rel,
                    source=self.name,
                    url="",
                    updated_at=mtime,
                    metadata={"path": rel, "ext": ext, "mode": "documents"},
                )
